"""
generate_sprint1_doc.py
Generates Sprint 1 – PortFolioPH implementation documentation as a Word .docx file.
Output: docs/Sprint1_PortFolioPH_Implementation_Report.docx
"""

import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colour constants ───────────────────────────────────────────────────────────
PRIMARY   = RGBColor(0x0D, 0x47, 0xA1)   # Deep Blue
ACCENT    = RGBColor(0xFF, 0x98, 0x00)   # Orange
DARK_GREY = RGBColor(0x42, 0x42, 0x42)
MID_GREY  = RGBColor(0x75, 0x75, 0x75)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GREEN     = RGBColor(0x38, 0x8E, 0x3C)
LIGHT_BLUE_BG = RGBColor(0xE3, 0xF2, 0xFD)

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set a table cell's background shading colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_run_color(run, rgb: RGBColor):
    run.font.color.rgb = rgb


def add_heading(doc: Document, text: str, level: int,
                color: RGBColor = PRIMARY, space_before: int = 12):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(space_before)
    for run in p.runs:
        run.font.color.rgb = color
    return p


def add_para(doc: Document, text: str, bold: bool = False,
             italic: bool = False, color: RGBColor = DARK_GREY,
             size: int = 11, space_after: int = 6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    set_run_color(run, color)
    return p


def add_bullet(doc: Document, text: str,
               color: RGBColor = DARK_GREY, size: int = 11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(size)
    set_run_color(run, color)
    return p


def make_table(doc: Document, headers: list, rows: list,
               header_bg: str = '0D47A1', col_widths: list = None):
    """Create a styled table with coloured header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr = table.rows[0]
    for idx, h in enumerate(headers):
        cell = hdr.cells[idx]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row in enumerate(rows):
        tr = table.rows[r_idx + 1]
        bg = 'E3F2FD' if r_idx % 2 == 0 else 'FFFFFF'
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(10)
            set_run_color(run, DARK_GREY)

    # Column widths
    if col_widths:
        for row in table.rows:
            for idx, w in enumerate(col_widths):
                row.cells[idx].width = Inches(w)

    return table


# ══════════════════════════════════════════════════════════════════════════════
# Document build
# ══════════════════════════════════════════════════════════════════════════════

def build_document() -> Document:
    doc = Document()

    # ── Page margins ───────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ══════════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ══════════════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    title = doc.add_heading('PortFolioPH', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = PRIMARY
        run.font.size = Pt(36)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Sprint 1 – Core Setup & Architecture\nImplementation Report')
    r.font.size = Pt(16)
    r.font.color.rgb = ACCENT
    r.bold = True

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lines = [
        ('Developer:  ', 'Mark Leannie Gacutno'),
        ('Project:    ', 'PortFolioPH – Offline-first Portfolio Builder'),
        ('Sprint:     ', 'Sprint 1 (Week 1)'),
        ('Date:       ', 'March 5, 2026'),
        ('Story Points:', '32'),
        ('Status:     ', '✅ COMPLETE'),
    ]
    for label, value in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_label = p.add_run(label)
        r_label.bold = True
        r_label.font.size = Pt(12)
        r_label.font.color.rgb = PRIMARY
        r_value = p.add_run(value)
        r_value.font.size = Pt(12)
        r_value.font.color.rgb = DARK_GREY

    doc.add_page_break()

    # ══════════════════════════════════════════════════════════════════════════
    # 1. PROJECT OVERVIEW
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '1. Project Overview', 1)
    add_para(doc, (
        'PortFolioPH is an offline-first Flutter Android application designed to allow '
        'students and fresh graduates to create, manage, and showcase professional '
        'digital portfolios directly from their mobile device — no internet connection required.'
    ))

    make_table(doc,
        headers=['Property', 'Value'],
        rows=[
            ('Application Name',  'PortFolioPH'),
            ('Platform',          'Android (API 26 – 34)'),
            ('Architecture',      'Clean Architecture + Provider (ChangeNotifier)'),
            ('State Management',  'Provider only'),
            ('Database',          'SQLite via sqflite (offline-first)'),
            ('Routing',           'GoRouter 14+'),
            ('Theme',             'Material 3, light + dark, primary #0D47A1'),
            ('Sprint',            'Sprint 1 – Core Setup & Architecture'),
            ('Story Points',      '32 hours'),
        ],
        col_widths=[2.0, 4.5],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 2. EPIC & STORIES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '2. Epic & User Stories', 1)
    add_para(doc, 'Epic ID: EPIC-001 | Label: sprint-1, foundation, flutter-setup | Fix Version: Sprint 1 (Week 1)')
    doc.add_paragraph()

    make_table(doc,
        headers=['Story ID', 'Title', 'Status', 'Effort'],
        rows=[
            ('STORY-001', 'Flutter Project Initialization & Package Setup',      '✅ Done', '2 hrs'),
            ('STORY-002', 'Clean Architecture Folder Structure',                  '✅ Done', '3 hrs'),
            ('STORY-003', 'SQLite DatabaseService & Schema (10 tables)',          '✅ Done', '5 hrs'),
            ('STORY-004', 'AppConstants & AppTheme (Material 3)',                 '✅ Done', '3 hrs'),
            ('STORY-005', 'GoRouter Setup with All Named Routes',                 '✅ Done', '4 hrs'),
            ('STORY-006', 'Bottom Navigation Scaffold + Placeholder Tabs (5)',    '✅ Done', '4 hrs'),
            ('STORY-007', 'Splash Screen with Session Check',                     '✅ Done', '4 hrs'),
            ('STORY-008', 'GitHub Setup, Android Permissions & Integration Test', '✅ Done', '4 hrs'),
            ('STORY-009', 'Sprint 1 Documentation & Handover',                    '✅ Done', '3 hrs'),
        ],
        col_widths=[1.2, 3.6, 1.0, 0.8],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 3. IMPLEMENTED FILES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '3. Implemented Files', 1)

    sections_files = [
        ('3.1  pubspec.yaml – Dependency Manifest', [
            'Added 13 production + 3 dev dependencies.',
            'sqflite ^2.3.3+1 – local SQLite engine.',
            'provider ^6.1.2 – state management via ChangeNotifier.',
            'go_router ^14.3.0 – declarative routing with named routes.',
            'shared_preferences ^2.3.3 – session persistence (userId).',
            'permission_handler ^11.3.1 – runtime Android permissions.',
            'crypto ^3.0.5 – SHA-256 password hashing.',
            'path_provider ^2.1.4, path ^1.9.0 – DB file path resolution.',
            'image_picker ^1.1.2, cached_network_image ^3.4.1 – media.',
            'intl ^0.20.2, uuid ^4.5.1, flutter_svg ^2.0.10+1 – utilities.',
            'Assets declared: assets/images/, assets/icons/.',
        ]),
        ('3.2  lib/core/constants/app_constants.dart', [
            'Abstract final class – cannot be instantiated.',
            'Defines: appName, appVersion, appTagline.',
            'DB constants: dbName = "portfolioph.db", dbVersion = 1.',
            'SharedPreferences keys: prefUserId, prefThemeMode, prefOnboardingDone.',
            'Brand colours: primaryColor #0D47A1, accentColor #FF9800.',
            'Typography scale: fontSizeXs(10) → fontSizeDisplay(32).',
            'Spacing: spacingXs(4)→spacingXxl(48), radii, elevations.',
            'Animation durations: durationFast(150ms), splashDuration(3s).',
            'BottomNav index constants: navIndexHome=0 … navIndexProfile=4.',
            'Validation limits: maxUsernameLength, minPasswordLength, etc.',
        ]),
        ('3.3  lib/core/theme/app_theme.dart', [
            'Two static getters: AppTheme.light, AppTheme.dark.',
            'Material 3 (useMaterial3: true) with ColorScheme.fromSeed.',
            'Light: primary #0D47A1, secondary #FF9800, surface #F5F5F5.',
            'Dark: primary #90CAF9 (accessible), surface #121212.',
            'Styled: AppBarTheme, BottomNavigationBarTheme, ElevatedButtonTheme,',
            '        TextButtonTheme, InputDecorationTheme, CardThemeData, DividerTheme.',
            '_buildTextTheme() shared builder for both modes.',
            'Zero magic numbers – every value from AppConstants.',
        ]),
        ('3.4  lib/core/router/app_router.dart', [
            'AppRoutes abstract final class – all route path constants.',
            'Routes: /, /login, /register, /dashboard.',
            'Future placeholders: /portfolio/new, /portfolio/:id, /project/new,',
            '  /project/:id, /resume/education/new, /resume/experience/new, /settings.',
            'AppRouter.create(UserProvider) factory → returns GoRouter.',
            'Auth redirect guard: unauthenticated users → /login; authenticated on auth routes → /dashboard.',
            'Splash route always passes through (manages its own redirect).',
            'debugLogDiagnostics: true enabled for development.',
        ]),
        ('3.5  lib/core/utils/helpers.dart', [
            'hashPassword(String) – SHA-256 hex string.',
            'formatDate(isoDate, pattern) – locale-safe date formatting.',
            'toIsoString(DateTime) / nowIso() – UTC ISO-8601 helpers.',
            'toTitleCase(String) – word capitalisation.',
            'initials(String fullName) – extracts "MG" from "Mark Gacutno".',
            'isValidEmail(String) – regex validation.',
            'isValidUrl(String) – Uri.parse validation.',
        ]),
        ('3.6  lib/data/datasources/local/database_service.dart', [
            'Singleton via factory constructor (DatabaseService._internal()).',
            'getDatabase() – lazy-open; open() – explicit open; close() – explicit close.',
            '_onConfigure() – enables PRAGMA foreign_keys = ON on every connection.',
            '_onCreate() – calls _runMigration1 on first install.',
            '_onUpgrade() – switch/case migration framework for v2+.',
            '_runMigration1() – db.batch() creates all 10 tables + 7 indexes atomically.',
            'All columns typed: INTEGER, TEXT, NOT NULL, DEFAULT, UNIQUE, FOREIGN KEY … ON DELETE CASCADE.',
        ]),
    ]

    for title_text, bullets in sections_files:
        add_heading(doc, title_text, 2, space_before=8)
        for b in bullets:
            add_bullet(doc, b)
        doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 4. DATABASE SCHEMA
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '4. Database Schema (10 Tables)', 1)
    add_para(doc, 'DB name: portfolioph.db | Version: 1 | FK enforcement: ON | All timestamps: ISO-8601 TEXT')
    doc.add_paragraph()

    schema_rows = [
        ('#', 'Table Name',        'Primary Key', 'FK References',     'Notable Columns'),
        ('1', 'users',             'id (AI)',      '—',                  'username UNIQUE, email UNIQUE, password_hash'),
        ('2', 'portfolios',        'id (AI)',      'users(id) CASCADE',  'is_public INT, template_id, custom_url'),
        ('3', 'projects',          'id (AI)',      'portfolios + users', 'tech_stack, is_featured, sort_order, thumbnail_path'),
        ('4', 'skills',            'id (AI)',      'users(id) CASCADE',  'category, level TEXT, years_of_experience'),
        ('5', 'education',         'id (AI)',      'users(id) CASCADE',  'institution, degree, field_of_study, is_current'),
        ('6', 'work_experience',   'id (AI)',      'users(id) CASCADE',  'company, job_title, employment_type, is_current'),
        ('7', 'certifications',    'id (AI)',      'users(id) CASCADE',  'issuing_organization, credential_id, does_expire'),
        ('8', 'contacts',          'id (AI)',      'users(id) CASCADE',  'platform, url, display_label, sort_order'),
        ('9', 'theme_settings',    'id (AI)',      'users(id) CASCADE',  'theme_mode TEXT, primary_color_hex | UNIQUE user_id'),
        ('10','app_settings',      'id (AI)',      'users(id) CASCADE',  'setting_key, setting_value | UNIQUE(user_id, key)'),
    ]

    table = doc.add_table(rows=11, cols=5)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for r_idx, row in enumerate(schema_rows):
        tr = table.rows[r_idx]
        is_header = r_idx == 0
        bg = '0D47A1' if is_header else ('E3F2FD' if r_idx % 2 == 1 else 'FFFFFF')
        for c_idx, cell_text in enumerate(row):
            cell = tr.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.size = Pt(9)
            run.bold = is_header
            run.font.color.rgb = WHITE if is_header else DARK_GREY

    for row in table.rows:
        widths = [0.3, 1.3, 1.0, 1.5, 2.5]
        for idx, w in enumerate(widths):
            row.cells[idx].width = Inches(w)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 5. MODELS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '5. Data Models', 1)
    add_para(doc, 'All 10 models follow the same contract: fromMap(Map<String, dynamic>), toMap(), copyWith(). Booleans stored as INTEGER 0/1 in SQLite.')

    make_table(doc,
        headers=['Model File', 'Class', 'Enum/Extension'],
        rows=[
            ('user_model.dart',          'UserModel',            '—'),
            ('portfolio_model.dart',     'PortfolioModel',       '—'),
            ('project_model.dart',       'ProjectModel',         '—'),
            ('skill_model.dart',         'SkillModel',           'SkillLevel (beginner…expert), SkillLevelExtension'),
            ('education_model.dart',     'EducationModel',       '—'),
            ('experience_model.dart',    'ExperienceModel',      '—'),
            ('certification_model.dart', 'CertificationModel',   '—'),
            ('contact_model.dart',       'ContactModel',         '—'),
            ('theme_setting_model.dart', 'ThemeSettingModel',    'AppThemeMode (light/dark/system), AppThemeModeExtension'),
            ('app_setting_model.dart',   'AppSettingModel',      '—'),
        ],
        col_widths=[2.2, 2.0, 2.4],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 6. REPOSITORIES
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '6. Repositories', 1)
    add_para(doc, 'Each repository receives DatabaseService via constructor injection (defaults to singleton). All SQL is parameterised — no string concatenation.')

    make_table(doc,
        headers=['Repository', 'Key Methods'],
        rows=[
            ('UserRepository',          'insert, findById, findByEmail, findByUsername, authenticate (SHA-256 compare), update, delete'),
            ('PortfolioRepository',     'insert, findById, findByUserId, update, delete'),
            ('ProjectRepository',       'insert, findById, findByPortfolioId, findFeaturedByUserId, update, delete'),
            ('SkillRepository',         'insert, findByUserId, findByCategory, update, delete'),
            ('EducationRepository',     'insert, findByUserId, update, delete'),
            ('ExperienceRepository',    'insert, findByUserId, update, delete'),
            ('CertificationRepository', 'insert, findByUserId, update, delete'),
            ('ContactRepository',       'insert, findByUserId, update, delete'),
        ],
        col_widths=[2.0, 4.6],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 7. PROVIDERS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '7. Providers (State Management)', 1)

    provider_data = [
        ('UserProvider', [
            'Manages currentUser (UserModel?) and isAuthenticated.',
            'restoreSession() – reads userId from SharedPreferences, loads from DB.',
            'login(email, password) – calls UserRepository.authenticate(); persists session.',
            'logout() – clears SharedPreferences, nulls currentUser, notifies listeners.',
            'updateProfile(UserModel) – updates DB and local state.',
            'Exposes isLoading, errorMessage; clearError() for UI.',
        ]),
        ('ThemeProvider', [
            'Holds ThemeMode (_themeMode), persisted in SharedPreferences key "themeMode".',
            'load() – called in main() before runApp() to prevent theme flicker.',
            'setThemeMode(mode) – update + persist.',
            'toggleDarkMode() – convenience toggle.',
        ]),
        ('NavigationProvider', [
            'Holds _currentIndex (0–4), matching AppConstants.navIndex* values.',
            'goTo(int) – guards against redundant rebuilds.',
            'Convenience: goHome(), goPortfolio(), goResume(), goSkills(), goProfile().',
        ]),
        ('PortfolioProvider', [
            'Holds _portfolios (List<PortfolioModel>) and _featuredProjects.',
            'loadForUser(int userId) – parallel load portfolios + featured projects.',
            'addPortfolio(PortfolioModel) – insert to DB + update local list.',
            'Full CRUD for projects/portfolios planned for Sprint 3.',
        ]),
    ]

    for p_name, bullets in provider_data:
        add_heading(doc, p_name, 2, space_before=8)
        for b in bullets:
            add_bullet(doc, b)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 8. SCREENS & WIDGETS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '8. Screens & Widgets', 1)

    make_table(doc,
        headers=['Screen / Widget', 'File', 'Status', 'Sprint'],
        rows=[
            ('SplashScreen',         'splash/splash_screen.dart',             '✅ Full',        'Sprint 1'),
            ('LoginScreen',          'auth/login_screen.dart',                '✅ Shell',       'Sprint 1 (full Sprint 2)'),
            ('RegisterScreen',       'auth/register_screen.dart',             '✅ Shell',       'Sprint 1 (full Sprint 2)'),
            ('MainScaffold',         'main_scaffold.dart',                    '✅ Full',        'Sprint 1'),
            ('DashboardScreen',      'dashboard/dashboard_screen.dart',       '🔜 Placeholder', 'Sprint 3'),
            ('PortfolioScreen',      'portfolio/portfolio_screen.dart',       '🔜 Placeholder', 'Sprint 3'),
            ('ResumeScreen',         'resume/resume_screen.dart',             '🔜 Placeholder', 'Sprint 4'),
            ('SkillsScreen',         'skills/skills_screen.dart',             '🔜 Placeholder', 'Sprint 4'),
            ('ProfileScreen',        'profile/profile_screen.dart',           '🔜 Placeholder', 'Sprint 5'),
            ('PlaceholderTabBody',   'widgets/common/placeholder_tab_body',   '✅ Full',        'Sprint 1'),
            ('LoadingWidget',        'widgets/common/loading_widget',         '✅ Full',        'Sprint 1'),
            ('AppErrorWidget',       'widgets/common/app_error_widget',       '✅ Full',        'Sprint 1'),
        ],
        col_widths=[2.0, 2.8, 1.2, 0.8],
    )
    doc.add_paragraph()

    add_heading(doc, 'SplashScreen – Detailed Flow', 2, space_before=8)
    steps = [
        '1. AnimationController (600ms) fades in logo + tagline.',
        '2. WidgetsBinding.addPostFrameCallback → calls _init() after first paint.',
        '3. Future.wait([DatabaseService().open(), Future.delayed(3s)]) – concurrent.',
        '4. context.read<UserProvider>().restoreSession() – queries SharedPreferences + DB.',
        '5. hasSession == true  → context.go("/dashboard")',
        '   hasSession == false → context.go("/login")',
    ]
    for s in steps:
        add_bullet(doc, s)

    add_heading(doc, 'MainScaffold – Tab Architecture', 2, space_before=8)
    add_para(doc, 'Uses IndexedStack (5 children) so tab state (scroll position, loaded data) is preserved on tab switch. BottomNavigationBar is driven entirely by NavigationProvider; no setState() in the scaffold widget itself.')
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 9. ROUTING
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '9. Routing (GoRouter)', 1)

    make_table(doc,
        headers=['Path', 'Name', 'Screen', 'Guard'],
        rows=[
            ('/',                       'splash',              'SplashScreen',    'None (always passes)'),
            ('/login',                  'login',               'LoginScreen',     'Redirects auth users → /dashboard'),
            ('/register',              'register',            'RegisterScreen',   'Redirects auth users → /dashboard'),
            ('/dashboard',             'dashboard',           'MainScaffold',     'Redirects unauth  → /login'),
            ('/portfolio/new',         'portfolio-new',       '(Sprint 3)',       'Protected'),
            ('/portfolio/:id',         'portfolio-detail',    '(Sprint 3)',       'Protected'),
            ('/project/new',           'project-new',         '(Sprint 3)',       'Protected'),
            ('/project/:id',           'project-detail',      '(Sprint 3)',       'Protected'),
            ('/resume/education/new',  'resume-education-new','(Sprint 4)',       'Protected'),
            ('/resume/experience/new', 'resume-experience-new','(Sprint 4)',      'Protected'),
            ('/settings',              'settings',            '(Sprint 6)',       'Protected'),
        ],
        col_widths=[1.8, 1.8, 1.4, 1.6],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 10. ANDROID PERMISSIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '10. Android Permissions (AndroidManifest.xml)', 1)

    make_table(doc,
        headers=['Permission', 'Purpose', 'API Range'],
        rows=[
            ('CAMERA',                'Profile photo & project screenshot capture',   'All'),
            ('READ_EXTERNAL_STORAGE', 'Read media files (legacy storage model)',       'API ≤ 32'),
            ('WRITE_EXTERNAL_STORAGE','Write files (resume PDF, assets)',              'API ≤ 29'),
            ('READ_MEDIA_IMAGES',     'Read images (granular media permission)',       'API ≥ 33'),
            ('INTERNET',             'Future sync + cached_network_image CDN',        'All'),
        ],
        col_widths=[2.2, 3.0, 1.4],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 11. ARCHITECTURE DIAGRAM
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '11. Architecture Diagram (Text)', 1)

    arch = doc.add_paragraph()
    arch.paragraph_format.space_after = Pt(4)
    arch_run = arch.add_run(
        '┌───────────────────────────────────────────────────┐\n'
        '│                  PRESENTATION LAYER                │\n'
        '│  Screens  ◄──  Providers (ChangeNotifier)          │\n'
        '│  GoRouter (AppRouter)  ◄──  UserProvider (guard)   │\n'
        '└─────────────────────┬─────────────────────────────┘\n'
        '                      │  reads / writes\n'
        '┌─────────────────────▼─────────────────────────────┐\n'
        '│                    DATA LAYER                       │\n'
        '│  Repositories  ──►  DatabaseService (Singleton)    │\n'
        '│                         │                           │\n'
        '│                    SQLite DB                        │\n'
        '│               (10 tables, FK ON, v1)               │\n'
        '└─────────────────────────────────────────────────────┘\n'
        '\n'
        'Session Storage:  SharedPreferences (userId, themeMode)\n'
        'Auth hashing:     crypto (SHA-256)\n'
        'Navigation state: NavigationProvider → IndexedStack (5 tabs)\n'
    )
    arch_run.font.name = 'Courier New'
    arch_run.font.size = Pt(9)
    arch_run.font.color.rgb = DARK_GREY
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 12. TECH STACK SUMMARY
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '12. Tech Stack Summary', 1)

    make_table(doc,
        headers=['Category', 'Package', 'Version', 'Usage'],
        rows=[
            ('State',        'provider',               '^6.1.2',    'ChangeNotifier-based state across app'),
            ('Routing',      'go_router',              '^14.3.0',   'Named routes, auth guard, deep link ready'),
            ('Database',     'sqflite',                '^2.3.3+1',  'SQLite offline-first storage'),
            ('Database',     'path_provider',          '^2.1.4',    'Resolves Documents dir for DB file'),
            ('Database',     'path',                   '^1.9.0',    'join() for DB file path'),
            ('Session',      'shared_preferences',     '^2.3.3',    'userId + themeMode persistence'),
            ('Permissions',  'permission_handler',     '^11.3.1',   'Runtime camera / storage requests'),
            ('Security',     'crypto',                 '^3.0.5',    'SHA-256 password hashing'),
            ('Media',        'image_picker',           '^1.1.2',    'Camera/gallery image selection'),
            ('Media',        'cached_network_image',   '^3.4.1',    'CDN image caching (future sprint)'),
            ('Utilities',    'intl',                   '^0.20.2',   'Date formatting, localisation'),
            ('Utilities',    'uuid',                   '^4.5.1',    'UUID generation for records'),
            ('UI',           'flutter_svg',            '^2.0.10+1', 'SVG icon rendering'),
        ],
        col_widths=[1.0, 1.8, 1.0, 2.8],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 13. SPRINT ROADMAP
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '13. Sprint Roadmap', 1)

    make_table(doc,
        headers=['Sprint', 'Focus Area', 'Status'],
        rows=[
            ('Sprint 1',  'Core Setup & Architecture (this report)',       '✅ Complete'),
            ('Sprint 2',  'Authentication – Login, Register, Session',      '🔜 Next'),
            ('Sprint 3',  'Portfolio & Projects CRUD + Project Detail',     '🔜'),
            ('Sprint 4',  'Resume – Education, Experience, Certifications', '🔜'),
            ('Sprint 5',  'Skills Management + Category Filters',           '🔜'),
            ('Sprint 6',  'Profile Edit + Settings Screen',                 '🔜'),
            ('Sprint 7',  'Export (PDF Resume) + Sharing',                  '🔜'),
            ('Sprint 8',  'Polish, Testing, CI/CD & Release APK',           '🔜'),
        ],
        col_widths=[1.0, 3.8, 1.6],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 14. DEFINITION OF DONE
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '14. Definition of Done – Sprint 1', 1)

    dod_items = [
        ('✅', 'pubspec.yaml contains all Sprint 1 dependencies; flutter pub get = 0 errors.'),
        ('✅', 'lib/ folder follows exact Clean Architecture scaffold.'),
        ('✅', 'DatabaseService opens SQLite, creates 10 tables with FK enforcement.'),
        ('✅', 'AppConstants – zero magic numbers or hardcoded strings in any other file.'),
        ('✅', 'AppTheme – Material 3, light + dark modes, brand colours.'),
        ('✅', 'AppRouter – GoRouter with auth guard, all named routes defined.'),
        ('✅', 'main.dart – MultiProvider + MaterialApp.router wired correctly.'),
        ('✅', 'Splash screen fades in, opens DB, checks session, redirects.'),
        ('✅', 'BottomNavigationBar – 5 tabs with IndexedStack (state preserved).'),
        ('✅', 'AndroidManifest.xml – CAMERA, READ/WRITE_EXTERNAL_STORAGE, READ_MEDIA_IMAGES.'),
        ('✅', 'flutter analyze – 0 errors, 0 warnings.'),
        ('✅', 'README.md updated with architecture, schema table, sprint roadmap.'),
    ]

    for icon, text in dod_items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_icon = p.add_run(icon + '  ')
        r_icon.font.size = Pt(11)
        r_icon.font.color.rgb = GREEN
        r_text = p.add_run(text)
        r_text.font.size = Pt(11)
        r_text.font.color.rgb = DARK_GREY

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 15. RISKS & MITIGATIONS
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '15. Risks & Mitigations', 1)

    make_table(doc,
        headers=['Risk', 'Impact', 'Mitigation', 'Status'],
        rows=[
            ('Flutter/Android SDK version mismatch',
             'High',
             'Ran flutter doctor; locked SDK range API 26–34 in build.gradle',
             '✅ Resolved'),
            ('Schema changes after Sprint 1',
             'Medium',
             'ERD finalized and signed off before coding; migration framework in onUpgrade() ready from v2',
             '✅ Mitigated'),
            ('GoRouter learning curve',
             'Low',
             'Named routes + AppRoutes constants isolate routing changes to one file',
             '✅ Resolved'),
            ('Permission denial on older Android',
             'Medium',
             'permission_handler with maxSdkVersion guards; READ_MEDIA_IMAGES for API 33+',
             '✅ Mitigated'),
        ],
        col_widths=[1.8, 0.7, 3.0, 1.0],
    )
    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════════════════
    # 16. SIGN-OFF
    # ══════════════════════════════════════════════════════════════════════════
    add_heading(doc, '16. Sprint 1 Sign-Off', 1)

    make_table(doc,
        headers=['Role', 'Name', 'Signature / Status', 'Date'],
        rows=[
            ('Developer', 'Mark Leannie Gacutno', '✅ Submitted', 'March 5, 2026'),
            ('Reviewer',  'Tom (Team Lead)',       '☐ Pending',   '____________'),
            ('Reviewer',  'Rex (QA Lead)',         '☐ Pending',   '____________'),
        ],
        col_widths=[1.2, 2.0, 2.0, 1.4],
    )
    doc.add_paragraph()
    add_para(doc,
        'Note: Team sign-off is required before Sprint 2 (Authentication) begins.',
        italic=True, color=MID_GREY)

    # ── Footer note ────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph()
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(6)
    r_hr = hr.add_run('─' * 80)
    r_hr.font.size = Pt(9)
    r_hr.font.color.rgb = MID_GREY

    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_footer = footer_p.add_run(
        'PortFolioPH  |  Sprint 1 Implementation Report  |  Prepared by Mark Leannie Gacutno  |  March 5, 2026'
    )
    r_footer.font.size = Pt(9)
    r_footer.font.color.rgb = MID_GREY
    r_footer.italic = True

    return doc


# ── Entry point ───────────────────────────────────────────────────────────────
if __name__ == '__main__':
    out_dir = os.path.join(os.path.dirname(__file__), '..', 'docs')
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, 'Sprint1_PortFolioPH_Implementation_Report.docx')

    document = build_document()
    document.save(out_path)

    abs_path = os.path.abspath(out_path)
    print(f'✅ Document saved:  {abs_path}')
    print(f'   Pages:          ~16')
    print(f'   Sections:        16')
    print(f'   Tables:          12')
